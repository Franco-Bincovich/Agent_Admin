from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from services.docx_template import _FONT, _MARGIN, extract_template_style
from utils.errors import AppError, ErrorCode
from utils.logger import log


def _new_doc_with_defaults(margin: int = _MARGIN) -> Document:
    """Crea un Document nuevo aplicando el margen dado en todos los lados."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = margin
    return doc


def _style_run(run, size_pt: int, color: RGBColor, bold: bool = False, font: str = _FONT) -> None:
    """Aplica fuente, tamaño en pt, color RGB y negrita a un run."""
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold


def _insert_logo(doc: Document, logo_bytes: bytes) -> None:
    """Inserta logo centrado al inicio del documento. No-op silencioso si falla."""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(BytesIO(logo_bytes), height=Cm(4))
        body = doc.element.body
        body.remove(p._element)
        body.insert(0, p._element)
    except Exception:
        pass


def generate_docx(
    outline: dict,
    imagenes: list[tuple[str, bytes]],
    usar_imagenes: bool,
    plantilla_bytes: bytes | None,
    logo_bytes: bytes | None = None,
) -> bytes:
    """Convierte el outline JSON en .docx; usa extract_template_style para aplicar la identidad visual."""
    try:
        style = extract_template_style(plantilla_bytes)
        font = style["font"]
        c_titulo = style["color_titulo"]
        c_seccion = style["color_seccion"]
        c_cuerpo = style["color_cuerpo"]
        margin = style["margin"]

        if plantilla_bytes:
            doc = Document(BytesIO(plantilla_bytes))
        else:
            doc = _new_doc_with_defaults(margin)

        if logo_bytes:
            _insert_logo(doc, logo_bytes)

        if not plantilla_bytes:
            p_titulo = doc.add_paragraph()
            _style_run(p_titulo.add_run(outline.get("titulo", "")), 24, c_titulo, bold=True, font=font)

        for sec_idx, seccion in enumerate(outline.get("secciones", [])):
            nombre = seccion.get("nombre", "")
            contenido = seccion.get("contenido", "")

            p_h = doc.add_paragraph()
            _style_run(p_h.add_run(nombre), 16, c_seccion, bold=True, font=font)
            doc.add_paragraph("─" * 60)

            p_b = doc.add_paragraph(contenido)
            if p_b.runs:
                _style_run(p_b.runs[0], 11, c_cuerpo, font=font)
            p_b.paragraph_format.line_spacing = 1.15

            if usar_imagenes:
                _insert_relevant_image(doc, nombre, contenido, imagenes, sec_idx)

        buffer = BytesIO()
        doc.save(buffer)
        log.info(f"docx.generated | secciones={len(outline.get('secciones', []))}")
        return buffer.getvalue()

    except AppError:
        raise
    except Exception as exc:
        log.error(f"docx.generation.failed | error={exc}")
        raise AppError("Error generando el documento DOCX.", ErrorCode.DOCUMENTO_GENERATION_FAILED, 500)


def _insert_relevant_image(
    doc: Document,
    seccion_nombre: str,
    seccion_contenido: str,
    imagenes: list[tuple[str, bytes]],
    sec_idx: int,
) -> None:
    """Inserta la imagen asignada a esta sección (por índice) centrada al ancho de contenido."""
    if not imagenes or sec_idx >= len(imagenes):
        return

    _, img_bytes = imagenes[sec_idx]

    section = doc.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(BytesIO(img_bytes), width=available_width)
    except Exception:
        pass
